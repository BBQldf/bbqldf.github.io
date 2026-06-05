#!/usr/bin/env python3
"""Convert docx files in ../self into Markdown drafts.

Images are extracted into a separate PicGotest checkout and Markdown image
references point at raw.githubusercontent.com URLs.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import unicodedata
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable
from urllib.parse import quote

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


SCRIPT_DIR = Path(__file__).resolve().parent
BLOG_REPO = SCRIPT_DIR.parent
WORKSPACE_ROOT = BLOG_REPO.parent


@dataclass
class ConversionStats:
    source: str
    slug: str
    output: str
    asset_dir: str
    paragraphs: int = 0
    tables: int = 0
    images: int = 0


def slugify(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    normalized = normalized.lower()
    normalized = re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "-", normalized)
    normalized = re.sub(r"-+", "-", normalized).strip("-")
    if normalized:
        return normalized[:80].strip("-")
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()[:10]
    return f"doc-{digest}"


def yaml_quote(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def guess_title(docx_path: Path, document: DocxDocument) -> str:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            return text
    return docx_path.stem


def iter_blocks(document: DocxDocument) -> Iterable[Paragraph | Table]:
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def markdown_escape_table_cell(value: str) -> str:
    value = re.sub(r"\s+", " ", value.strip())
    return value.replace("|", r"\|")


def table_to_markdown(table: Table) -> list[str]:
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([markdown_escape_table_cell(cell.text) for cell in row.cells])
    if not rows:
        return []

    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    header = normalized[0]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    for row in normalized[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return lines


def image_extension(partname: str) -> str:
    suffix = Path(partname).suffix.lower()
    return suffix if suffix else ".png"


def write_image(part, asset_dir: Path, image_index: int) -> tuple[str, str]:
    ext = image_extension(str(part.partname))
    filename = f"image-{image_index:02d}{ext}"
    target = asset_dir / filename
    target.write_bytes(part.blob)
    return filename, ext


def paragraph_to_markdown(
    paragraph: Paragraph,
    document: DocxDocument,
    asset_dir: Path,
    image_base_url: str,
    image_index: int,
) -> tuple[list[str], int, int]:
    lines: list[str] = []
    text_chunks: list[str] = []
    image_count = 0

    for run in paragraph.runs:
        if run.text:
            text_chunks.append(run.text)

        blips = run._element.xpath(".//*[local-name()='blip']")
        for blip in blips:
            rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rel_id:
                continue
            part = document.part.related_parts.get(rel_id)
            if part is None:
                continue

            pending_text = "".join(text_chunks).strip()
            if pending_text:
                lines.append(pending_text)
                text_chunks = []

            image_index += 1
            image_count += 1
            filename, _ = write_image(part, asset_dir, image_index)
            lines.append(f"![]({image_base_url}/{filename})")

    text = "".join(text_chunks).strip()
    if text:
        style = paragraph.style.name if paragraph.style is not None else ""
        if style.startswith("Heading"):
            level_match = re.search(r"(\d+)", style)
            level = int(level_match.group(1)) if level_match else 2
            level = min(max(level, 2), 6)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    return lines, image_index, image_count


def convert_docx(
    docx_path: Path,
    output_dir: Path,
    pic_repo: Path,
    asset_prefix: str,
    raw_base_url: str,
    date: str,
) -> ConversionStats:
    document = Document(str(docx_path))
    title = guess_title(docx_path, document)
    slug = slugify(docx_path.stem)

    asset_dir = pic_repo / asset_prefix / slug
    if asset_dir.exists():
        shutil.rmtree(asset_dir)
    asset_dir.mkdir(parents=True, exist_ok=True)

    encoded_slug = quote(slug)
    image_base_url = f"{raw_base_url.rstrip('/')}/{encoded_slug}"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{slug}.md"

    stats = ConversionStats(
        source=str(docx_path),
        slug=slug,
        output=str(output_path),
        asset_dir=str(asset_dir),
    )

    lines: list[str] = [
        "---",
        "layout:     post",
        f"title:      {yaml_quote(title)}",
        f"subtitle:   {yaml_quote('Word 转换草稿，待人工校对')}",
        f"date:       {date}",
        "author:     bbq",
        "header-img: img/post-bg-digital-native.jpg",
        "catalog: true",
        "project: self",
        "tags:",
        "    - self",
        "    - Word转换",
        "---",
        "",
        f"# {title}",
        "",
        f"> Source: `{docx_path.name}`",
        "",
    ]

    image_index = 0
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            para_lines, image_index, images_added = paragraph_to_markdown(
                block,
                document,
                asset_dir,
                image_base_url,
                image_index,
            )
            if para_lines:
                stats.paragraphs += 1
                lines.extend(para_lines)
                lines.append("")
            stats.images += images_added
        elif isinstance(block, Table):
            table_lines = table_to_markdown(block)
            if table_lines:
                stats.tables += 1
                lines.extend(table_lines)
                lines.append("")

    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")

    if stats.images == 0 and asset_dir.exists() and not any(asset_dir.iterdir()):
        asset_dir.rmdir()

    return stats


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source-root",
        type=Path,
        default=WORKSPACE_ROOT / "self",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=WORKSPACE_ROOT / "self" / "converted",
    )
    parser.add_argument(
        "--pic-repo",
        type=Path,
        default=WORKSPACE_ROOT / "PicGotest",
    )
    parser.add_argument("--asset-prefix", default="blog-content/self")
    parser.add_argument(
        "--raw-base-url",
        default="https://raw.githubusercontent.com/BBQldf/PicGotest/master/blog-content/self",
    )
    parser.add_argument("--date", default="2026-06-05")
    args = parser.parse_args()

    docx_files = sorted(args.source_root.rglob("*.docx"))
    if not docx_files:
        raise SystemExit(f"No docx files found under {args.source_root}")
    if not args.pic_repo.exists():
        raise SystemExit(f"PicGotest repo not found: {args.pic_repo}")

    stats = [
        convert_docx(
            docx_path,
            args.output_dir,
            args.pic_repo,
            args.asset_prefix,
            args.raw_base_url,
            args.date,
        )
        for docx_path in docx_files
    ]

    manifest_path = args.output_dir / "manifest.json"
    manifest_path.write_text(
        json.dumps([stat.__dict__ for stat in stats], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    total_images = sum(stat.images for stat in stats)
    print(f"Converted {len(stats)} docx files")
    print(f"Extracted {total_images} images")
    print(f"Markdown output: {args.output_dir}")
    print(f"Image output: {args.pic_repo / args.asset_prefix}")
    print(f"Manifest: {manifest_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
